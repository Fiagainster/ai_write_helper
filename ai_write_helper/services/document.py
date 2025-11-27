#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""文档服务模块

该模块实现文档的读取和写入功能，支持不同格式的文档处理。
"""

import logging
import os
import tempfile
from threading import RLock


class DocumentService:
    """文档服务类，负责文档的读取和写入操作"""
    
    def __init__(self, config_manager):
        """初始化文档服务
        
        Args:
            config_manager: 配置管理器实例
        """
        self.logger = logging.getLogger("ai_write_helper.document")
        self.lock = RLock()  # 可重入锁，用于线程安全
        self.config_manager = config_manager
    
    def validate_path(self, file_path):
        """验证文件路径是否有效
        
        Args:
            file_path: 文件路径
            
        Returns:
            bool: 如果路径有效且可访问返回True，否则返回False
        """
        try:
            # 检查路径是否存在
            dir_path = os.path.dirname(file_path)
            if dir_path and not os.path.exists(dir_path):
                self.logger.warning(f"目录不存在: {dir_path}")
                return False
            
            # 如果文件已存在，检查是否有读写权限
            if os.path.exists(file_path):
                if not os.access(file_path, os.R_OK | os.W_OK):
                    self.logger.warning(f"无文件读写权限: {file_path}")
                    return False
            else:
                # 如果文件不存在，检查目录是否有写入权限
                if dir_path and not os.access(dir_path, os.W_OK):
                    self.logger.warning(f"无目录写入权限: {dir_path}")
                    return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"验证文件路径时出错: {str(e)}")
            return False
    
    def read_document(self, file_path):
        """读取文档内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文档内容
        """
        with self.lock:
            try:
                # 首先验证文件是否存在且可访问
                if not self.validate_path(file_path):
                    self.logger.error(f"无效的文件路径: {file_path}")
                    return ""
                
                # 检查文件大小，避免读取空文件
                if os.path.getsize(file_path) == 0:
                    self.logger.warning(f"文件为空: {file_path}")
                    return ""
                
                # 获取文件扩展名
                _, ext = os.path.splitext(file_path)
                ext = ext.lower()
                
                # 根据文件类型选择不同的读取方法
                if ext in ['.txt', '.md']:
                    return self._read_text_file(file_path)
                elif ext == '.docx':
                    # 先尝试使用文本方式读取，检查是否是有效的文本文件
                    try:
                        # 检查文件前几个字节，判断是否是有效的docx文件
                        with open(file_path, 'rb') as f:
                            header = f.read(4)
                            # docx文件应该是zip格式，以PK开头
                            if header != b'PK\x03\x04':
                                self.logger.warning(f"文件 {file_path} 扩展名是.docx，但不是有效的zip文件，尝试以文本方式读取")
                                return self._read_text_file(file_path)
                    except Exception as e:
                        self.logger.warning(f"检查docx文件格式时出错: {str(e)}，尝试以文本方式读取")
                        return self._read_text_file(file_path)
                    
                    # 如果是有效的docx文件，使用专门的方法读取
                    return self._read_docx_file(file_path)
                else:
                    # 默认为文本文件
                    self.logger.warning(f"不支持的文件格式: {ext}，尝试以文本方式读取")
                    return self._read_text_file(file_path)
                    
            except Exception as e:
                self.logger.error(f"读取文档时出错: {str(e)}")
                # 出错时返回空字符串，允许程序继续执行
                return ""
    
    def write_document(self, file_path, content, incremental=False):
        """写入文档内容（默认为重写模式）
        
        Args:
            file_path: 文件路径
            content: 要写入的内容
            incremental: 是否增量更新（默认为False，即完整重写）
        """
        with self.lock:
            # 验证文件路径
            if not self.validate_path(file_path):
                raise ValueError(f"无效的文件路径: {file_path}")
                
            # 获取文件扩展名
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # 记录写入模式
            mode = "重写" if not incremental else "增量"
            self.logger.info(f"开始{mode}写入文档: {file_path}")
            self.logger.debug(f"写入内容长度: {len(content)} 字符")
            
            # 原子性写入 - 先写入临时文件，再替换
            temp_dir = os.path.dirname(file_path)
            if not temp_dir:
                temp_dir = os.getcwd()
            
            # 创建临时文件路径
            temp_path = tempfile.mktemp(dir=temp_dir)
            
            try:
                # 根据文件类型选择不同的写入方法
                if ext in ['.txt', '.md']:
                    self._write_text_file(temp_path, file_path, content, incremental)
                elif ext == '.docx':
                    self._write_docx_file(temp_path, file_path, content, incremental)
                else:
                    # 默认为文本文件
                    self.logger.warning(f"不支持的文件格式: {ext}，尝试以文本方式写入")
                    self._write_text_file(temp_path, file_path, content, incremental)
                
                self.logger.info(f"文档{mode}写入成功: {file_path}")
                    
            except Exception as e:
                self.logger.error(f"写入文档时出错: {str(e)}")
                # 添加详细错误信息以便调试
                import traceback
                self.logger.debug(f"错误详情: {traceback.format_exc()}")
                raise
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    try:
                        os.unlink(temp_path)
                    except Exception as e:
                        self.logger.error(f"清理临时文件时出错: {str(e)}")
    
    def _read_text_file(self, file_path):
        """读取文本文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文件内容
        """
        # 尝试不同的编码格式
        encodings = ['utf-8', 'gbk', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                self.logger.info(f"使用编码 {encoding} 读取文件: {file_path}")
                return content
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，抛出异常
        raise UnicodeDecodeError(
            'utf-8',
            b'',
            0,
            1,
            '无法使用任何支持的编码格式解码文件'
        )
    
    def _write_text_file(self, temp_path, target_path, content, incremental):
        """写入文本文件，增强重写模式
        
        Args:
            temp_path: 临时文件路径
            target_path: 目标文件路径
            content: 要写入的内容
            incremental: 是否增量更新
        """
        # 确保内容是字符串类型
        if not isinstance(content, str):
            content = str(content)
        
        # 确保内容不为空
        content = content.strip()
        if not content:
            self.logger.warning("写入内容为空")
        
        # 使用原子写入方式，确保文件完整性
        if not incremental:
            # 对于重写模式，使用临时文件进行原子替换，避免写入中断导致文件损坏
            self.logger.debug(f"使用原子写入方式重写文件: {target_path}")
            
            # 写入内容到临时文件
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(content)
                # 确保内容被刷新到磁盘
                f.flush()
                os.fsync(f.fileno())
            
            # 原子性替换文件
            self._atomic_replace(temp_path, target_path)
            self.logger.debug(f"原子替换操作成功完成")
        else:
            # 增量模式
            # 读取现有内容
            existing_content = ''
            if os.path.exists(target_path):
                existing_content = self._read_text_file(target_path)
            
            # 写入内容到临时文件
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(existing_content)
                # 在现有内容和新内容之间添加空行，使内容更加清晰
                if existing_content and content:
                    if not existing_content.endswith('\n'):
                        f.write('\n')
                    f.write('\n')
                f.write(content)
                # 确保内容被刷新到磁盘
                f.flush()
                os.fsync(f.fileno())
            
            # 原子性替换文件
            self._atomic_replace(temp_path, target_path)
    
    def _read_docx_file(self, file_path):
        """读取Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文档内容
        """
        try:
            # 尝试导入python-docx库
            import docx
            
            # 确保文件存在且可访问
            if not os.path.exists(file_path):
                self.logger.error(f"文件不存在: {file_path}")
                # 尝试使用备用方案
                return self._extract_docx_text(file_path)
            
            if not os.access(file_path, os.R_OK):
                self.logger.error(f"无文件读取权限: {file_path}")
                # 尝试使用备用方案
                return self._extract_docx_text(file_path)
            
            # 尝试使用python-docx读取
            self.logger.debug(f"使用python-docx读取Word文档: {file_path}")
            doc = docx.Document(file_path)
            content = []
            
            for para in doc.paragraphs:
                content.append(para.text)
            
            result = '\n'.join(content)
            self.logger.debug(f"成功使用python-docx读取Word文档，获取到 {len(result)} 字符")
            return result
            
        except ImportError:
            self.logger.warning("python-docx库未安装，尝试提取纯文本")
            # 作为备用方案，可以将docx作为zip文件处理
            return self._extract_docx_text(file_path)
        except Exception as e:
            self.logger.error(f"使用python-docx读取Word文档时出错: {str(e)}")
            self.logger.info("尝试使用备用方案提取纯文本")
            # 尝试使用备用方案
            return self._extract_docx_text(file_path)
    
    def _write_docx_file(self, temp_path, target_path, content, incremental):
        """写入Word文档，增强重写模式
        
        Args:
            temp_path: 临时文件路径
            target_path: 目标文件路径
            content: 要写入的内容
            incremental: 是否增量更新
        """
        try:
            import docx
            
            # 确保内容是字符串类型
            if not isinstance(content, str):
                content = str(content)
            
            content = content.strip()
            if not content:
                self.logger.warning("写入内容为空")
            
            # 处理文件写入
            if incremental and os.path.exists(target_path):
                # 增量模式：打开现有文档
                try:
                    doc = docx.Document(target_path)
                    # 添加新内容，按空行分割段落
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                except Exception as e:
                    self.logger.error(f"读取现有文档失败: {str(e)}")
                    # 如果读取失败，则创建新文档
                    doc = docx.Document()
                    # 按空行分割段落
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
            else:
                # 完整重写模式
                self.logger.debug(f"执行Word文档完整重写: {target_path}")
                doc = docx.Document()
                # 对于重写模式，按段落分割内容
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            # 保存到临时文件
            temp_docx_path = temp_path + '.docx'
            doc.save(temp_docx_path)
            
            # 原子性替换文件
            self._atomic_replace(temp_docx_path, target_path)
            self.logger.debug(f"Word文档原子替换操作成功完成")
            
        except ImportError:
            self.logger.warning("python-docx库未安装，将内容保存为文本文件")
            # 作为备用方案，保存为文本文件
            self._write_text_file(temp_path, target_path + '.txt', content, incremental)
        except Exception as e:
            self.logger.error(f"写入Word文档时出错: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path):
        """从docx文件中提取纯文本（备用方案）
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 提取的文本
        """
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            # docx文件本质上是一个zip文件
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # 读取document.xml文件
                with zip_ref.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    # namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # 提取文本
                    text_parts = []
                    for para in root.findall('.//w:p', ns):
                        para_text = []
                        for text_elem in para.findall('.//w:t', ns):
                            if text_elem.text:
                                para_text.append(text_elem.text)
                        if para_text:
                            text_parts.append(''.join(para_text))
                    
                    return '\n'.join(text_parts)
                    
        except Exception as e:
            self.logger.error(f"提取docx文本时出错: {str(e)}")
            # 如果提取失败，返回空字符串
            return ""
    
    def _atomic_replace(self, source_path, target_path):
        """原子性替换文件，增强版本
        
        Args:
            source_path: 源文件路径
            target_path: 目标文件路径
        """
        import shutil
        import stat
        
        # 确保目标文件的目录存在
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        try:
            # 在Windows上使用shutil.move实现原子替换
            if os.name == 'nt':
                # 先删除目标文件（如果存在）
                if os.path.exists(target_path):
                    # 确保文件可写
                    try:
                        os.chmod(target_path, stat.S_IWRITE)
                    except:
                        self.logger.warning("无法修改目标文件权限")
                    # 删除文件
                    os.remove(target_path)
                # 移动源文件到目标位置
                shutil.move(source_path, target_path)
            else:
                # 在类Unix系统上使用os.replace
                os.replace(source_path, target_path)
                
            self.logger.debug(f"成功将文件从 {source_path} 原子替换到 {target_path}")
            self.logger.info(f"成功写入文件: {target_path}")
            
        except Exception as e:
            # 如果os.replace失败，降级使用shutil.copy2
            self.logger.warning(f"原子替换失败，尝试使用非原子性方法: {str(e)}")
            import shutil
            shutil.copy2(source_path, target_path)
            self.logger.warning(f"使用非原子性方法复制文件")